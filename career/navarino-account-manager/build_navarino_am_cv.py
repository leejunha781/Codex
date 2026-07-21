#!/usr/bin/env python3
"""Build NAVARINO Account Manager revised CV and Korean match strategy docs.

Canonical Windows paths (NOT overwritten; not available in this cloud workspace):
  e:\\이력서\\Account Manager\\Job Ad_Account Manager.docx
  e:\\이력서\\Account Manager\\Joonha_Lee_NAVARINO_Account_Manager_Professional_CV.docx
  e:\\이력서\\Account Manager\\Navarino_한국_Account_Manager_회사조사_JD매칭_합격전략_20260719.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor, Twips


OUT_DIR = Path(__file__).resolve().parent
CV_PATH = OUT_DIR / "Joonha_Lee_NAVARINO_Account_Manager_Professional_CV_Revised_JD_Fit_20260721.docx"
MATCH_PATH = OUT_DIR / "Navarino_한국_Account_Manager_회사조사_JD매칭_합격전략_20260721_Revised.docx"
README_PATH = OUT_DIR / "README.md"

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT = RGBColor(0x2C, 0x5F, 0x8A)
GRAY = RGBColor(0x44, 0x44, 0x44)
RULE = RGBColor(0x2C, 0x5F, 0x8A)


def set_run_font(run, name: str = "Calibri", size: int = 10, bold: bool = False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2C5F8A")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_narrow_margins(doc: Document):
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)


def tight_paragraph(paragraph, before: float = 0, after: float = 4, line: float = 1.08):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_heading_bar(doc: Document, text: str):
    p = doc.add_paragraph()
    tight_paragraph(p, before=10, after=2)
    run = p.add_run(text.upper())
    set_run_font(run, size=11, bold=True, color=NAVY)
    add_horizontal_line(p)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    set_run_font(run, size=9.5, color=GRAY)
    tight_paragraph(p, before=0, after=2, line=1.05)
    return p


def add_job_header(doc: Document, title: str, company: str, dates: str, location: str):
    p = doc.add_paragraph()
    tight_paragraph(p, before=8, after=0)
    r1 = p.add_run(title)
    set_run_font(r1, size=10.5, bold=True, color=NAVY)
    r2 = p.add_run(f"  |  {company}")
    set_run_font(r2, size=10, bold=True, color=ACCENT)

    p2 = doc.add_paragraph()
    tight_paragraph(p2, before=0, after=3)
    r3 = p2.add_run(f"{dates}  ·  {location}")
    set_run_font(r3, size=9, color=GRAY)
    return p


def build_cv() -> Path:
    doc = Document()
    set_narrow_margins(doc)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Header
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tight_paragraph(name, before=0, after=2)
    nr = name.add_run("JOONHA LEE")
    set_run_font(nr, size=20, bold=True, color=NAVY)

    role = doc.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tight_paragraph(role, before=0, after=2)
    rr = role.add_run("Account Manager  ·  Maritime Connectivity, IT & Digital Solutions")
    set_run_font(rr, size=11, bold=True, color=ACCENT)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tight_paragraph(contact, before=0, after=2)
    cr = contact.add_run(
        "Gunpo-si, Gyeonggi-do, Republic of Korea  ·  "
        "leejunha781@gmail.com  ·  "
        "linkedin.com/in/joonha-lee-20b518316"
    )
    set_run_font(cr, size=9, color=GRAY)
    add_horizontal_line(contact)

    # Professional Summary
    add_heading_bar(doc, "Professional Summary")
    summary = (
        "Account-oriented Engineering and Business Development professional with 21+ years in "
        "maritime, naval shipbuilding, defence electronics, and satellite communications. "
        "Combines customer-facing technical engagement—system integration, verification, FAT/SAT, "
        "and commissioning—with stakeholder alignment suited to long-cycle B2B account care. "
        "Professional English communicator with KPI/quality-gate discipline and conceptual fluency "
        "in commercial PLM, SAP/ERP, and maritime digital operations (self-directed PLM study while "
        "preparing for AVEVA Consultant applications—not employed AVEVA delivery). Seeking Navarino "
        "Account Manager (Korea) to grow relationships, support pre- and post-sale experience, and "
        "build commercial pipeline ownership with integrity."
    )
    sp = doc.add_paragraph()
    tight_paragraph(sp, before=2, after=4, line=1.1)
    sr = sp.add_run(summary)
    set_run_font(sr, size=9.5, color=GRAY)

    # Core Value Proposition
    add_heading_bar(doc, "Selected Outcomes Aligned to Account Management")
    for item in [
        "Built durable customer trust through end-to-end maritime system delivery—from interface design and integration testing to FAT/SAT and commissioning handovers.",
        "Translated complex satcom and shipboard communication capabilities into clear acceptance criteria and stakeholder-aligned decisions for demanding maritime customers.",
        "Practised regulatory, traceability, and quality-control discipline across naval programmes—transferable to commercial fleet compliance, SLA, and performance-management conversations.",
        "Operated as a reliable cross-functional interface for pre-sale clarity and post-sale validation—practising retention/expansion behaviours and commercial pipeline discipline for a full Account Manager remit.",
    ]:
        add_bullet(doc, item)

    # Skills
    add_heading_bar(doc, "Core Competencies")

    skills = [
        (
            "Commercial & Account Care",
            "Consultative customer engagement · Relationship building & retention mindset · "
            "Pre-sales / post-sales technical support · Stakeholder alignment · Brand representation · "
            "Pipeline & activity discipline (CRM-ready) · Commercial proposal support · Professional English",
        ),
        (
            "Maritime Domain",
            "Naval / maritime programmes · Shipboard communication & entertainment systems · "
            "Satellite connectivity (LEO / satcom) · Hybrid network & bandwidth operational awareness · "
            "FAT / SAT / commissioning · Integration & verification · Regulatory documentation & traceability",
        ),
        (
            "Digital / SaaS-Adjacent Fluency",
            "Maritime digital transformation narrative · Commercial PLM concepts (self-directed learning "
            "initiated while preparing for AVEVA Consultant applications) · SAP/ERP familiarity · "
            "KPI / SLA / performance frameworks · QC & acceptance culture · Conversation fluency for "
            "onboard IT, RMM/asset inventory, cybersecurity/OT, and IT-as-a-Service solution themes",
        ),
        (
            "Tools & Working Methods",
            "Jira · Confluence · ERP/PLM concepts · Interface control & wiring documentation · "
            "Project scheduling & QA gates · Evidence-based commissioning records",
        ),
    ]
    for title, body in skills:
        p = doc.add_paragraph()
        tight_paragraph(p, before=3, after=1)
        t = p.add_run(f"{title}: ")
        set_run_font(t, size=9.5, bold=True, color=NAVY)
        b = p.add_run(body)
        set_run_font(b, size=9.5, color=GRAY)

    # Experience
    add_heading_bar(doc, "Professional Experience")

    add_job_header(
        doc,
        "General Manager",
        "Intellian Technologies",
        "Mar 2023 – Jul 2025",
        "Seongnam, Gyeonggi, South Korea",
    )
    for b in [
        "Oversaw design, integration, and verification of LEO satellite terminal systems, aligning technical performance with customer acceptance expectations in a regulated maritime connectivity context.",
        "Managed embedded-system integration for LEO parabolic and AESA-related architectures, improving operational readiness and reducing ambiguity at customer validation milestones.",
        "Led project scheduling and quality-assurance gates for cross-functional delivery and customer-facing technical reviews; supported FAT/SAT and commissioning with clear stakeholder communication.",
        "Acted as a professional company interface to customers and partners, reinforcing brand trust through disciplined delivery, documentation, and issue ownership.",
        "Applied KPI-minded progress tracking and Jira/Confluence-style collaboration habits to keep priorities visible across engineering and customer stakeholders.",
    ]:
        add_bullet(doc, b)

    add_job_header(
        doc,
        "Senior System Engineer",
        "GENOHCO Inc.",
        "Aug 2022 – Mar 2023",
        "Gunpo-si, Gyeonggi, South Korea",
    )
    for b in [
        "Contributed to the K2 Sight Maintenance Equipment System, practising defence-grade quality expectations, configuration discipline, and stakeholder communication.",
        "Supported Unmanned Aerial Vehicle (UAV) Maintenance Equipment programmes across hardware-oriented system-engineering workstreams.",
        "Developed PCB solutions for UAV maintenance equipment, reinforcing detail orientation, documentation accuracy, and schedule reliability—foundations for account hygiene and customer confidence.",
    ]:
        add_bullet(doc, b)

    add_job_header(
        doc,
        "Senior Manager",
        "Daeyang Electric Co., Ltd.",
        "May 2005 – Apr 2022",
        "Busan, South Korea",
    )
    for b in [
        "Delivered multi-year naval communication and entertainment programmes (Indonesia submarine Entertainment System; PC Communication & Broadcasting Systems), managing interface clarity, integration outcomes, and customer satisfaction across long programme lifecycles.",
        "Spearheaded ICS development for a 300-ton class submarine and supported 3000-ton class submarine ICS; developed internal wireless/broadcasting for Korean frigates under demanding technical and regulatory constraints.",
        "Performed V/UHF and satcom signal specification analysis, wiring diagrams, integration testing, and interface-box design for military secure equipment—building evidence-first, cybersecurity/OT-aware discipline transferable to commercial maritime SLA conversations.",
        "Sustained long-horizon relationships with shipyards, technical managers, and programme stakeholders—demonstrating the retention and trust behaviours expected of a maritime Account Manager.",
    ]:
        add_bullet(doc, b)

    # Education
    add_heading_bar(doc, "Education")
    edu1 = doc.add_paragraph()
    tight_paragraph(edu1, before=4, after=0)
    e1 = edu1.add_run("M.S., Embedded Systems / Electrical & Electronics Engineering")
    set_run_font(e1, size=10, bold=True, color=NAVY)
    edu1b = doc.add_paragraph()
    tight_paragraph(edu1b, before=0, after=2)
    e1b = edu1b.add_run("Pusan National University  ·  2013 – 2017  ·  GPA 3.79  ·  South Korea")
    set_run_font(e1b, size=9, color=GRAY)
    add_bullet(
        doc,
        "Focus: embedded systems, hardware/PCB design, industrial networking, OS-level programming (Linux/Windows); research on real-time embedded control with industrial Ethernet (EtherCAT).",
    )

    edu2 = doc.add_paragraph()
    tight_paragraph(edu2, before=6, after=0)
    e2 = edu2.add_run("B.S., Control & Instrumentation Engineering")
    set_run_font(e2, size=10, bold=True, color=NAVY)
    edu2b = doc.add_paragraph()
    tight_paragraph(edu2b, before=0, after=2)
    e2b = edu2b.add_run("University of Ulsan  ·  South Korea")
    set_run_font(e2b, size=9, color=GRAY)

    # Why Navarino
    add_heading_bar(doc, "Alignment with Navarino (Company Fit Narrative)")
    for b in [
        "Vision fit: represent Navarino’s ambition to be the benchmark maritime technology company—Integrity, Excellence, and Innovation in every customer interaction.",
        "Portfolio conversation fluency (learning, not product ownership): Starlink/hybrid connectivity, Infinity bandwidth/network management, Spectrum RMM/asset inventory, Quazar IT-as-a-Service, cybersecurity/OT (Angel/Ozora/NEMO).",
        "Commercial motion: retention and expansion via consultative discovery with operators, technical managers, IT, and compliance-adjacent roles; growing fluency with DPA/crewing decision chains; CRM hygiene via PLM/ERP/Jira/Confluence/KPI literacy.",
    ]:
        add_bullet(doc, b)

    # Languages
    add_heading_bar(doc, "Languages")
    lp = doc.add_paragraph()
    tight_paragraph(lp, before=2, after=2)
    lr = lp.add_run(
        "Korean (Native)  ·  English (Professional — customer engagement, written proposals, cross-border coordination)"
    )
    set_run_font(lr, size=9.5, color=GRAY)

    doc.save(CV_PATH)
    return CV_PATH


def build_match_doc() -> Path:
    doc = Document()
    set_narrow_margins(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tight_paragraph(title, before=0, after=4)
    tr = title.add_run("NAVARINO Account Manager (Korea)\n회사조사 · JD 매칭 · 합격 전략 (2026-07-21 Revised)")
    set_run_font(tr, size=14, bold=True, color=NAVY)

    meta = doc.add_paragraph()
    tight_paragraph(meta, before=0, after=6)
    mr = meta.add_run(
        "기준 문서: Navarino_한국_Account_Manager_회사조사_JD매칭_합격전략_20260719 (원본 보존) · "
        "공개 JD / Navarino Who We Are · LinkedIn + MEMORY 경력 사실만 반영"
    )
    set_run_font(mr, size=9, color=GRAY)

    add_heading_bar(doc, "1. 회사 요약 (Vision / Mission / Portfolio)")
    for b in [
        "Vision: To be the benchmark technology company of the maritime world.",
        "Mission: Develop, distribute and support innovative technology services that add true value to maritime customers while contributing to a sustainable future.",
        "Scale: 650+ clients · 13,000+ vessels · 25+ years · offices across Greece, UK, Cyprus, Germany, Norway, Netherlands, Singapore, Hong Kong, Japan, UAE, USA, India (+ Korea hiring remote AM).",
        "Values: Integrity · Excellence · Innovation · Great Place to Work certified (multi-year).",
        "Core portfolio talking points: Infinity (bandwidth/network management smart-box), Spectrum (RMM/asset inventory), Quazar (IT-as-a-Service), cybersecurity (Angel/Ozora), connectivity (Starlink, Fleet Xpress, KU, hybrid), NEMO (OT entry), Navarino Elements (GMDSS/nav/safety electronics).",
    ]:
        add_bullet(doc, b)

    add_heading_bar(doc, "2. JD Must / Nice (Korea Account Manager)")
    for b in [
        "Must: related bachelor’s; maritime + regulatory environment; interpersonal + sales generation; brand ambassador; reliable/detail/prioritisation; Fluent English.",
        "Responsibilities: retention/satisfaction; cross-team pre/post sales; new business + strategy; CRM hygiene; ambassador; industry events.",
        "Nice: B2B / Maritime SaaS sales; operators / crewing managers / DPA understanding; QC systems & performance-management frameworks.",
    ]:
        add_bullet(doc, b)

    add_heading_bar(doc, "3. 매칭 점수 (Honest)")
    for b in [
        "BEFORE CV reframing: ~58% — 엔지니어링/프로그램 전달 중심 이력으로 읽힘.",
        "AFTER CV reframing (본 Revised JD Fit): ~74% — 동일 사실을 lifecycle account care + consultative solution selling + KPI/SLA/regulatory discipline + satcom/connectivity fluency로 재배치.",
        "80%+를 주장하려면: 상업용 merchant fleet 네트워크 증거, SaaS quota 실적, CRM tool 실무 경력을 추가로 확보해야 함 (날조 금지).",
    ]:
        add_bullet(doc, b)

    add_heading_bar(doc, "4. 갭과 보완 전략")
    for b in [
        "순수 상업 쿼터 세일즈: 커버레터·면접에서 30-60-90 (CRM 위생 → joint hunting → weekly pipeline)로 소유. 매출 숫자 창작 금지.",
        "한국 상선/선사 네트워크: 조선소·기술관리·satcom 바이어 경로를 owner/IT/procurement 맵으로 확장. 협회·행사 참가 의지 명시.",
        "Maritime SaaS quota: Infinity/Spectrum/Quazar 가치 제안 학습 + discovery question 세트로 증명. 가짜 케이스 스터디 금지.",
        "DPA/Crewing: ISM/DPA·crewing 의사결정 체인 학습. “규제된 선박 환경을 알고, commercial DPA 어휘를 빠르게 쌓겠다”로 정직하게 소유.",
        "PLM 서술: AVEVA Consultant 지원 준비 중 자가학습 시작 + 상업 PLM/SAP/ERP/Jira/Confluence/KPI 익숙 — employed AVEVA Marine 납품 경력으로 과장하지 않음. Marine SaaS와 유사점(구성관리, 변경 추적, 성과지표, 협업 툴)으로만 연결.",
    ]:
        add_bullet(doc, b)

    add_heading_bar(doc, "5. ATS 키워드 (사실 기반 weave)")
    for b in [
        "Account management · Customer relationship · Maritime / shipping · Satellite connectivity / LEO / satcom · FAT/SAT/commissioning · Pre-sales / post-sales · Stakeholder management · Regulatory / traceability · KPI / SLA · CRM discipline · B2B consultative selling · Cybersecurity / OT awareness",
    ]:
        add_bullet(doc, b)

    add_heading_bar(doc, "6. 클레임 금지 (면접 리스크)")
    for b in [
        "GENOHCO에서 satellite/TVAC 업무 주장 금지 (K2 Sight ME / UAV ME / PCB만).",
        "매출·쿼터%·ARR·폐쇄 로고 창작 금지.",
        "AVEVA Marine 고용 납품 / DPA·crewing을 기존 직무 이해관계자로 과장 금지.",
        "Navarino 650 clients / 13,000 vessels를 개인 성과처럼 서술 금지.",
    ]:
        add_bullet(doc, b)

    add_heading_bar(doc, "7. 산출물 경로")
    for b in [
        f"Revised CV: {CV_PATH.name}",
        "Preserved canonical (local Windows, not in cloud): e:\\이력서\\Account Manager\\Joonha_Lee_NAVARINO_Account_Manager_Professional_CV.docx",
        "Preserved strategy base: e:\\이력서\\Account Manager\\Navarino_한국_Account_Manager_회사조사_JD매칭_합격전략_20260719.docx",
        f"This revised strategy: {MATCH_PATH.name}",
    ]:
        add_bullet(doc, b)

    doc.save(MATCH_PATH)
    return MATCH_PATH


def write_readme(cv: Path, match: Path) -> None:
    README_PATH.write_text(
        f"""# NAVARINO Account Manager — Revised JD Fit (2026-07-21)

## Goal
Tailor Joonha Lee’s CV for **Navarino Account Manager (South Korea, remote + customer visits)** with formal English, LinkedIn/MEMORY-aligned facts, and honest match uplift.

## Source paths (canonical — preserved, not overwritten)
- `e:\\이력서\\Account Manager\\Job Ad_Account Manager.docx`
- `e:\\이력서\\Account Manager\\Joonha_Lee_NAVARINO_Account_Manager_Professional_CV.docx`
- `e:\\이력서\\Account Manager\\Navarino_한국_Account_Manager_회사조사_JD매칭_합격전략_20260719.docx`

> Note: Those Windows files were **not synced** into this Cursor Cloud workspace. Deliverables below are **new copies** built from public JD, Navarino company pages, LinkedIn, and MEMORY — not in-place edits of the canonical files.

## Outputs
| Artifact | Path |
|----------|------|
| Revised CV (EN) | `{cv.name}` |
| Match / strategy (KO+EN) | `{match.name}` |
| Builder | `build_navarino_am_cv.py` |

## Match score (honest)
- Before reframing: **~58%**
- After Revised JD Fit: **~74%**

## Role axis
**A — Sales / FAE / Solution (Account Manager)**. Marine PLM is framed as self-directed digital-ops fluency only (AVEVA application prep), not as employed Marine PLM delivery.

## Rebuild
```bash
python3 build_navarino_am_cv.py
```
""",
        encoding="utf-8",
    )


def main():
    cv = build_cv()
    match = build_match_doc()
    write_readme(cv, match)
    print(f"Wrote {cv}")
    print(f"Wrote {match}")
    print(f"Wrote {README_PATH}")


if __name__ == "__main__":
    main()
